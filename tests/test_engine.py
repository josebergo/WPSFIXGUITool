from __future__ import annotations

import io
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

from wpsfix.engine import convert_docx, inspect_docx, unique_output_path


def field_run(kind: str) -> OxmlElement:
    run = OxmlElement("w:r")
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), kind)
    run.append(node)
    return run


def instruction_run(text: str) -> OxmlElement:
    run = OxmlElement("w:r")
    node = OxmlElement("w:instrText")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


def make_image(path: Path, fmt: str, color: tuple[int, int, int]) -> None:
    Image.new("RGB", (180, 80), color).save(path, fmt)


def make_fixture(path: Path) -> None:
    work = path.parent
    png = work / "cached.png"
    tiff = work / "logo.tiff"
    make_image(png, "PNG", (224, 56, 48))
    make_image(tiff, "TIFF", (34, 112, 190))

    document = Document()
    header = document.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Inches(6.0))
    table.cell(0, 0).text = "页次"
    page_paragraph = table.cell(0, 1).paragraphs[0]._p
    page_paragraph.append(field_run("begin"))
    page_paragraph.append(instruction_run(" PAGE } / { NUMPAGES "))
    page_paragraph.append(field_run("separate"))
    cached = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    cached.append(text)
    page_paragraph.append(cached)
    page_paragraph.append(field_run("end"))
    header.add_paragraph().add_run().add_picture(str(tiff), width=Inches(1.2))

    paragraph = document.add_paragraph("缓存图片：")
    picture_run = paragraph.add_run()
    picture_run.add_picture(str(png), width=Inches(1.8))
    run_xml = picture_run._r
    paragraph_xml = paragraph._p
    paragraph_xml.remove(run_xml)
    paragraph_xml.append(field_run("begin"))
    paragraph_xml.append(instruction_run(r' INCLUDEPICTURE \\d "C:\\Temp\\missing.png" '))
    paragraph_xml.append(field_run("separate"))
    paragraph_xml.append(run_xml)
    paragraph_xml.append(field_run("end"))
    document.add_paragraph("正文内容。")
    document.save(path)


class EngineTests(unittest.TestCase):
    def test_unique_output_path_does_not_overwrite(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            source = Path(temp) / "demo.docx"
            source.write_bytes(b"x")
            first = unique_output_path(source)
            self.assertEqual(first.name, "demo-WPS兼容版.docx")
            first.write_bytes(b"x")
            self.assertEqual(unique_output_path(source).name, "demo-WPS兼容版-2.docx")

    def test_repairs_fields_and_cached_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            temp_path = Path(temp)
            source = temp_path / "fixture.docx"
            output = temp_path / "fixture-fixed.docx"
            make_fixture(source)

            before = inspect_docx(source)
            self.assertEqual(before.malformed_page_fields_found, 1)
            self.assertEqual(before.include_picture_fields_found, 1)

            progress: list[int] = []
            report = convert_docx(
                source, output,
                progress=lambda value, _: progress.append(value),
                update_with_word=False,
            )

            self.assertTrue(output.exists())
            self.assertTrue(output.with_suffix(".report.json").exists())
            self.assertEqual(report.page_fields_rebuilt, 1)
            self.assertEqual(report.include_picture_fields_unwrapped, 1)
            self.assertGreaterEqual(report.tiff_images_converted, 1)
            self.assertEqual(report.remaining_include_picture_fields, 0)
            self.assertTrue(report.zip_crc_ok)
            self.assertTrue(report.xml_parse_ok)
            self.assertTrue(report.image_relationships_ok)
            self.assertEqual(progress[-1], 100)

            with zipfile.ZipFile(output) as archive:
                combined = b"\n".join(archive.read(name) for name in archive.namelist() if name.endswith(".xml"))
                self.assertNotIn(b"INCLUDEPICTURE", combined)
                self.assertIn(b"NUMPAGES", combined)


if __name__ == "__main__":
    unittest.main()
